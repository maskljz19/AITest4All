"""Document Parser Service

Parses various document formats and extracts text content.
Supports: Word (.doc/.docx), PDF (.pdf), Markdown (.md), Excel (.xls/.xlsx), TXT (.txt), and URLs.
"""

import io
import logging
from typing import Optional
from pathlib import Path

import docx
import PyPDF2
import markdown
import openpyxl
import requests
from bs4 import BeautifulSoup

# Optional import for old Excel format
try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False
    logging.warning("xlrd not available. Old Excel format (.xls) support will be limited.")

logger = logging.getLogger(__name__)


class DocumentParseError(Exception):
    """Exception raised when document parsing fails"""
    pass


class DocumentParser:
    """Service for parsing various document formats"""
    
    SUPPORTED_EXTENSIONS = {'.doc', '.docx', '.pdf', '.md', '.xls', '.xlsx', '.txt'}
    REQUEST_TIMEOUT = 30  # seconds
    
    @classmethod
    def parse_file(cls, file_path: str) -> str:
        """Parse a file and extract text content
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentParseError: If parsing fails or format is unsupported
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentParseError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise DocumentParseError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )
        
        try:
            if extension in ['.doc', '.docx']:
                return cls._parse_docx(file_path)
            elif extension == '.pdf':
                return cls._parse_pdf(file_path)
            elif extension == '.md':
                return cls._parse_markdown(file_path)
            elif extension in ['.xls', '.xlsx']:
                return cls._parse_excel(file_path)
            elif extension == '.txt':
                return cls._parse_txt(file_path)
        except Exception as e:
            logger.error(f"Failed to parse file {file_path}: {str(e)}")
            raise DocumentParseError(f"Failed to parse file: {str(e)}")
    
    @classmethod
    def parse_url(cls, url: str) -> str:
        """Fetch and parse content from a URL
        
        Args:
            url: URL to fetch content from
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentParseError: If fetching or parsing fails
        """
        try:
            response = requests.get(url, timeout=cls.REQUEST_TIMEOUT)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '').lower()
            
            # Handle HTML content
            if 'text/html' in content_type:
                return cls._parse_html(response.text)
            
            # Handle plain text
            elif 'text/plain' in content_type:
                return response.text
            
            # Handle markdown
            elif 'text/markdown' in content_type:
                return cls._parse_markdown_text(response.text)
            
            # Default: try to parse as HTML
            else:
                return cls._parse_html(response.text)
                
        except requests.RequestException as e:
            logger.error(f"Failed to fetch URL {url}: {str(e)}")
            raise DocumentParseError(f"Failed to fetch URL: {str(e)}")
        except Exception as e:
            logger.error(f"Failed to parse URL content {url}: {str(e)}")
            raise DocumentParseError(f"Failed to parse URL content: {str(e)}")
    
    @classmethod
    def extract_text(cls, content: bytes, file_type: str) -> str:
        """Extract text from binary content
        
        Args:
            content: Binary content
            file_type: File type/extension (e.g., 'docx', 'pdf')
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentParseError: If extraction fails
        """
        file_type = file_type.lower().lstrip('.')
        
        try:
            if file_type in ['doc', 'docx']:
                return cls._parse_docx_bytes(content)
            elif file_type == 'pdf':
                return cls._parse_pdf_bytes(content)
            elif file_type == 'md' or file_type == 'markdown':
                return cls._parse_markdown_text(content.decode('utf-8'))
            elif file_type in ['xls', 'xlsx']:
                return cls._parse_excel_bytes(content)
            elif file_type == 'txt':
                return content.decode('utf-8')
            else:
                raise DocumentParseError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Failed to extract text from {file_type}: {str(e)}")
            raise DocumentParseError(f"Failed to extract text: {str(e)}")
    
    # Private parsing methods
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Parse Word document (.doc and .docx)
        
        Note: python-docx only supports .docx format.
        For .doc files, it will attempt to read them, but may fail.
        Consider using LibreOffice or other converters for .doc files.
        """
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return '\n\n'.join(paragraphs)
        except Exception as e:
            logger.error(f"Failed to parse Word document: {str(e)}")
            raise DocumentParseError(
                f"Failed to parse Word document. "
                f"Note: .doc format may not be fully supported. "
                f"Please convert to .docx format for best results. Error: {str(e)}"
            )
    
    @staticmethod
    def _parse_docx_bytes(content: bytes) -> str:
        """Parse Word document from bytes (.doc and .docx)"""
        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return '\n\n'.join(paragraphs)
        except Exception as e:
            logger.error(f"Failed to parse Word document from bytes: {str(e)}")
            raise DocumentParseError(
                f"Failed to parse Word document. "
                f"Note: .doc format may not be fully supported. "
                f"Please convert to .docx format for best results. Error: {str(e)}"
            )
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Parse PDF document"""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
    
    @staticmethod
    def _parse_pdf_bytes(content: bytes) -> str:
        """Parse PDF document from bytes"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text_parts = []
        
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _parse_markdown(file_path: str) -> str:
        """Parse Markdown document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_text = file.read()
            # Convert markdown to HTML then extract text
            html = markdown.markdown(md_text)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text(separator='\n\n', strip=True)
    
    @staticmethod
    def _parse_markdown_text(text: str) -> str:
        """Parse Markdown text"""
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text(separator='\n\n', strip=True)
    
    @staticmethod
    def _parse_excel(file_path: str) -> str:
        """Parse Excel document (.xls and .xlsx)"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        # Handle old .xls format
        if extension == '.xls':
            if not XLRD_AVAILABLE:
                raise DocumentParseError(
                    "xlrd library not available. Cannot parse .xls files. "
                    "Please convert to .xlsx format or install xlrd."
                )
            return DocumentParser._parse_xls(file_path)
        
        # Handle new .xlsx format
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"Sheet: {sheet_name}")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)
            
            text_parts.append('')  # Empty line between sheets
        
        return '\n'.join(text_parts)
    
    @staticmethod
    def _parse_xls(file_path: str) -> str:
        """Parse old Excel format (.xls)"""
        import xlrd
        
        workbook = xlrd.open_workbook(file_path)
        text_parts = []
        
        for sheet in workbook.sheets():
            text_parts.append(f"Sheet: {sheet.name}")
            
            for row_idx in range(sheet.nrows):
                row = sheet.row(row_idx)
                row_text = '\t'.join(str(cell.value) if cell.value else '' for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)
            
            text_parts.append('')  # Empty line between sheets
        
        return '\n'.join(text_parts)
    
    @staticmethod
    def _parse_excel_bytes(content: bytes) -> str:
        """Parse Excel document from bytes (.xls and .xlsx)"""
        # Try to detect format by checking file signature
        # .xls files start with D0 CF 11 E0 (OLE2 format)
        # .xlsx files start with 50 4B (ZIP format)
        
        if content[:2] == b'PK':  # ZIP format (.xlsx)
            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)
                
                text_parts.append('')
            
            return '\n'.join(text_parts)
        
        elif content[:8] == b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1':  # OLE2 format (.xls)
            if not XLRD_AVAILABLE:
                raise DocumentParseError(
                    "xlrd library not available. Cannot parse .xls files. "
                    "Please convert to .xlsx format or install xlrd."
                )
            
            import xlrd
            workbook = xlrd.open_workbook(file_contents=content)
            text_parts = []
            
            for sheet in workbook.sheets():
                text_parts.append(f"Sheet: {sheet.name}")
                
                for row_idx in range(sheet.nrows):
                    row = sheet.row(row_idx)
                    row_text = '\t'.join(str(cell.value) if cell.value else '' for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)
                
                text_parts.append('')
            
            return '\n'.join(text_parts)
        
        else:
            # Default to trying .xlsx format
            try:
                workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
                text_parts = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"Sheet: {sheet_name}")
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                        if row_text.strip():
                            text_parts.append(row_text)
                    
                    text_parts.append('')
                
                return '\n'.join(text_parts)
            except Exception as e:
                raise DocumentParseError(f"Failed to parse Excel file: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Parse plain text document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    @staticmethod
    def _parse_html(html_content: str) -> str:
        """Parse HTML content and extract text"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(['script', 'style']):
            script.decompose()
        
        # Get text
        text = soup.get_text(separator='\n\n', strip=True)
        
        # Clean up multiple newlines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n\n'.join(lines)
