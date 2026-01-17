"""Export API for Test Cases and Code"""

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse, Response
from typing import List, Optional
from pydantic import BaseModel
import json
import io
import zipfile
from datetime import datetime

from app.services.session_manager import SessionManager

router = APIRouter(prefix="/api/v1/export", tags=["export"])


# Dependency to get services
async def get_session_manager() -> SessionManager:
    """Get session manager instance"""
    from app.core.redis_client import get_redis
    redis = await get_redis()
    return SessionManager(redis)


# Request Models
class ExportCasesRequest(BaseModel):
    """Request model for exporting test cases"""
    session_id: str
    format: str  # excel/word/json/markdown/html/csv
    include_requirement: bool = False
    include_scenarios: bool = False
    include_cases: bool = True
    include_quality_report: bool = False


class ExportCodeRequest(BaseModel):
    """Request model for exporting code"""
    session_id: str
    format: str  # zip/single/project
    project_name: Optional[str] = "test_automation"


def generate_excel_export(data: dict) -> bytes:
    """Generate Excel file from test cases"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Test Cases"
        
        # Header style
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        
        # Headers
        headers = ["Case ID", "Title", "Type", "Priority", "Precondition", "Steps", "Expected Result"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Data
        cases = data.get("cases", [])
        for row, case in enumerate(cases, 2):
            ws.cell(row=row, column=1, value=case.get("case_id", ""))
            ws.cell(row=row, column=2, value=case.get("title", ""))
            ws.cell(row=row, column=3, value=case.get("test_type", ""))
            ws.cell(row=row, column=4, value=case.get("priority", ""))
            ws.cell(row=row, column=5, value=case.get("precondition", ""))
            
            # Format steps
            steps = case.get("steps", [])
            steps_text = "\n".join([
                f"{s.get('step_no', '')}. {s.get('action', '')}"
                for s in steps
            ])
            ws.cell(row=row, column=6, value=steps_text)
            ws.cell(row=row, column=7, value=case.get("expected_result", ""))
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 30
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 10
        ws.column_dimensions['E'].width = 30
        ws.column_dimensions['F'].width = 50
        ws.column_dimensions['G'].width = 30
        
        # Save to bytes
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output.getvalue()
        
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="openpyxl library not installed. Cannot generate Excel files."
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate Excel: {str(e)}"
        )


def generate_word_export(data: dict) -> bytes:
    """Generate Word document from test cases"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Test Cases Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        
        # Test Cases
        cases = data.get("cases", [])
        for i, case in enumerate(cases, 1):
            # Case title
            doc.add_heading(f"{i}. {case.get('title', 'Untitled')}", level=2)
            
            # Case details
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Fill table
            table.cell(0, 0).text = "Case ID"
            table.cell(0, 1).text = case.get("case_id", "")
            
            table.cell(1, 0).text = "Test Type"
            table.cell(1, 1).text = case.get("test_type", "")
            
            table.cell(2, 0).text = "Priority"
            table.cell(2, 1).text = case.get("priority", "")
            
            table.cell(3, 0).text = "Precondition"
            table.cell(3, 1).text = case.get("precondition", "")
            
            table.cell(4, 0).text = "Test Steps"
            steps = case.get("steps", [])
            steps_text = "\n".join([
                f"{s.get('step_no', '')}. {s.get('action', '')}\n   Data: {s.get('data', '')}\n   Expected: {s.get('expected', '')}"
                for s in steps
            ])
            table.cell(4, 1).text = steps_text
            
            table.cell(5, 0).text = "Expected Result"
            table.cell(5, 1).text = case.get("expected_result", "")
            
            table.cell(6, 0).text = "Postcondition"
            table.cell(6, 1).text = case.get("postcondition", "")
            
            doc.add_paragraph("")
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
        
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx library not installed. Cannot generate Word files."
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate Word document: {str(e)}"
        )


def generate_json_export(data: dict) -> bytes:
    """Generate JSON file from test cases"""
    try:
        json_str = json.dumps(data, indent=2, ensure_ascii=False)
        return json_str.encode('utf-8')
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate JSON: {str(e)}"
        )


def generate_markdown_export(data: dict) -> bytes:
    """Generate Markdown file from test cases"""
    try:
        lines = []
        lines.append("# Test Cases Document\n")
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        lines.append("---\n")
        
        cases = data.get("cases", [])
        for i, case in enumerate(cases, 1):
            lines.append(f"## {i}. {case.get('title', 'Untitled')}\n")
            lines.append(f"**Case ID:** {case.get('case_id', '')}\n")
            lines.append(f"**Test Type:** {case.get('test_type', '')}\n")
            lines.append(f"**Priority:** {case.get('priority', '')}\n")
            lines.append(f"**Precondition:** {case.get('precondition', '')}\n")
            
            lines.append("\n**Test Steps:**\n")
            steps = case.get("steps", [])
            for step in steps:
                lines.append(f"{step.get('step_no', '')}. {step.get('action', '')}\n")
                lines.append(f"   - Data: {step.get('data', '')}\n")
                lines.append(f"   - Expected: {step.get('expected', '')}\n")
            
            lines.append(f"\n**Expected Result:** {case.get('expected_result', '')}\n")
            lines.append(f"**Postcondition:** {case.get('postcondition', '')}\n")
            lines.append("\n---\n")
        
        return "".join(lines).encode('utf-8')
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate Markdown: {str(e)}"
        )


def generate_html_export(data: dict) -> bytes:
    """Generate HTML file from test cases"""
    try:
        html = []
        html.append("<!DOCTYPE html>")
        html.append("<html><head>")
        html.append("<meta charset='utf-8'>")
        html.append("<title>Test Cases Document</title>")
        html.append("<style>")
        html.append("body { font-family: Arial, sans-serif; margin: 20px; }")
        html.append("h1 { color: #333; }")
        html.append("h2 { color: #666; border-bottom: 2px solid #ddd; padding-bottom: 5px; }")
        html.append("table { border-collapse: collapse; width: 100%; margin: 10px 0; }")
        html.append("th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }")
        html.append("th { background-color: #366092; color: white; }")
        html.append(".step { margin: 5px 0; }")
        html.append("</style>")
        html.append("</head><body>")
        
        html.append("<h1>Test Cases Document</h1>")
        html.append(f"<p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")
        
        cases = data.get("cases", [])
        for i, case in enumerate(cases, 1):
            html.append(f"<h2>{i}. {case.get('title', 'Untitled')}</h2>")
            html.append("<table>")
            html.append(f"<tr><th>Case ID</th><td>{case.get('case_id', '')}</td></tr>")
            html.append(f"<tr><th>Test Type</th><td>{case.get('test_type', '')}</td></tr>")
            html.append(f"<tr><th>Priority</th><td>{case.get('priority', '')}</td></tr>")
            html.append(f"<tr><th>Precondition</th><td>{case.get('precondition', '')}</td></tr>")
            
            html.append("<tr><th>Test Steps</th><td>")
            steps = case.get("steps", [])
            for step in steps:
                html.append(f"<div class='step'>{step.get('step_no', '')}. {step.get('action', '')}</div>")
                html.append(f"<div style='margin-left: 20px;'>Data: {step.get('data', '')}</div>")
                html.append(f"<div style='margin-left: 20px;'>Expected: {step.get('expected', '')}</div>")
            html.append("</td></tr>")
            
            html.append(f"<tr><th>Expected Result</th><td>{case.get('expected_result', '')}</td></tr>")
            html.append(f"<tr><th>Postcondition</th><td>{case.get('postcondition', '')}</td></tr>")
            html.append("</table>")
        
        html.append("</body></html>")
        
        return "".join(html).encode('utf-8')
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate HTML: {str(e)}"
        )


def generate_csv_export(data: dict) -> bytes:
    """Generate CSV file from test cases"""
    try:
        import csv
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write header
        writer.writerow([
            "Case ID", "Title", "Test Type", "Priority", 
            "Precondition", "Steps", "Expected Result", "Postcondition"
        ])
        
        # Write data
        cases = data.get("cases", [])
        for case in cases:
            # Format steps
            steps = case.get("steps", [])
            steps_text = "; ".join([
                f"{s.get('step_no', '')}. {s.get('action', '')} (Data: {s.get('data', '')}, Expected: {s.get('expected', '')})"
                for s in steps
            ])
            
            writer.writerow([
                case.get("case_id", ""),
                case.get("title", ""),
                case.get("test_type", ""),
                case.get("priority", ""),
                case.get("precondition", ""),
                steps_text,
                case.get("expected_result", ""),
                case.get("postcondition", "")
            ])
        
        # Get CSV content
        output.seek(0)
        return output.getvalue().encode('utf-8-sig')  # Use utf-8-sig for Excel compatibility
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate CSV: {str(e)}"
        )


@router.post("/cases")
async def export_cases(
    request: ExportCasesRequest,
    session_manager: SessionManager = Depends(get_session_manager)
):
    """
    Export test cases in various formats
    
    Supports: Excel, Word, JSON, Markdown, HTML
    """
    try:
        # Get data from session
        cases = []
        scenarios = []
        requirement_analysis = None
        quality_report = None
        
        if request.include_cases:
            cases = await session_manager.get_step_result(request.session_id, "cases") or []
        
        if request.include_scenarios:
            scenarios = await session_manager.get_step_result(request.session_id, "scenarios") or []
        
        if request.include_requirement:
            requirement_analysis = await session_manager.get_step_result(request.session_id, "requirement_analysis")
        
        if request.include_quality_report:
            quality_report = await session_manager.get_step_result(request.session_id, "quality_report")
        
        # Check if we have data to export
        if not cases and not scenarios and not requirement_analysis and not quality_report:
            raise HTTPException(
                status_code=400,
                detail="No data available to export. Please generate test cases first."
            )
        
        # Prepare data
        data = {}
        
        if cases:
            data["cases"] = cases
        
        if scenarios:
            data["scenarios"] = scenarios
        
        if requirement_analysis:
            data["requirement_analysis"] = requirement_analysis
        
        if quality_report:
            data["quality_report"] = quality_report
        
        # Generate export based on format
        format_lower = request.format.lower()
        
        if format_lower == "excel":
            content = generate_excel_export(data)
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            filename = f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        elif format_lower == "word":
            content = generate_word_export(data)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        elif format_lower == "json":
            content = generate_json_export(data)
            media_type = "application/json"
            filename = f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        elif format_lower == "markdown":
            content = generate_markdown_export(data)
            media_type = "text/markdown"
            filename = f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        
        elif format_lower == "html":
            content = generate_html_export(data)
            media_type = "text/html"
            filename = f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        
        elif format_lower == "csv":
            content = generate_csv_export(data)
            media_type = "text/csv"
            filename = f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported format: {request.format}. Supported: excel, word, json, markdown, html, csv"
            )
        
        # Return file
        return Response(
            content=content,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to export cases: {str(e)}"
        )


@router.post("/code")
async def export_code(
    request: ExportCodeRequest,
    session_manager: SessionManager = Depends(get_session_manager)
):
    """
    Export test automation code
    
    Supports: ZIP (multiple files), single file, project structure
    """
    try:
        # Get code files from session
        code_result = await session_manager.get_step_result(request.session_id, "code")
        
        if not code_result:
            raise HTTPException(
                status_code=400,
                detail="No code available to export. Please generate code first."
            )
        
        # Extract files from result
        files = code_result.get("files", {}) if isinstance(code_result, dict) else {}
        
        if not files:
            raise HTTPException(
                status_code=400,
                detail="No code files found in session."
            )
        
        format_lower = request.format.lower()
        
        if format_lower == "zip":
            # Create ZIP file with all code files
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for filename, content in files.items():
                    zip_file.writestr(filename, content)
            
            zip_buffer.seek(0)
            content = zip_buffer.getvalue()
            media_type = "application/zip"
            filename = f"{request.project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        elif format_lower == "single":
            # Combine all files into a single file
            combined = []
            for filename, file_content in files.items():
                combined.append(f"# File: {filename}\n")
                combined.append(file_content)
                combined.append("\n\n")
            
            content = "".join(combined).encode('utf-8')
            media_type = "text/plain"
            filename = f"{request.project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        elif format_lower == "project":
            # Create project structure ZIP
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Add all files with proper directory structure
                for filepath, file_content in files.items():
                    zip_file.writestr(f"{request.project_name}/{filepath}", file_content)
                
                # Add README if not present
                if "README.md" not in files:
                    readme_content = f"""# {request.project_name}

Test automation project generated by AI Test Case Generator.

## Setup

1. Install dependencies:
   ```bash
   pip install -r requirements.txt
   ```

2. Run tests:
   ```bash
   pytest
   ```

Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
                    zip_file.writestr(f"{request.project_name}/README.md", readme_content)
            
            zip_buffer.seek(0)
            content = zip_buffer.getvalue()
            media_type = "application/zip"
            filename = f"{request.project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported format: {request.format}. Supported: zip, single, project"
            )
        
        # Return file
        return Response(
            content=content,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to export code: {str(e)}"
        )
